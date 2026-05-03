"""MR 일일/주간 보고서 API.

여러 메모(또는 일일 보고서)를 묶어 Gemini Flash 로 종합 정리한 결과를 저장/조회한다.
부족한 부분은 docx 다운로드 후 사용자가 직접 편집.
"""
import io
import json
import logging
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select, and_
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.auth.deps import get_current_user
from app.models.connection import get_db
from app.models.database import Doctor, Hospital, MemoTemplate, Report, User, VisitMemo
from app.schemas.schemas import ReportCreate, ReportResponse
from app.services.ai_memo import summarize_report

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/reports", tags=["보고서"])


def _parse_json_array(raw: Optional[str]) -> Optional[list[int]]:
    if not raw:
        return None
    try:
        v = json.loads(raw)
        return v if isinstance(v, list) else None
    except (TypeError, ValueError):
        return None


def _parse_ai(raw: Optional[str]):
    if not raw:
        return None
    try:
        return json.loads(raw)
    except (TypeError, ValueError):
        return raw


def _serialize(r: Report) -> dict:
    return {
        "id": r.id,
        "user_id": r.user_id,
        "report_type": r.report_type,
        "period_start": r.period_start,
        "period_end": r.period_end,
        "title": r.title,
        "source_memo_ids": _parse_json_array(r.source_memo_ids),
        "source_report_ids": _parse_json_array(r.source_report_ids),
        "raw_combined": r.raw_combined,
        "ai_summary": _parse_ai(r.ai_summary),
        "template_id": r.template_id,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
    }


def _ai_summary_to_text(ai_summary) -> str:
    """AI 정리 결과를 보고서 묶을 때 쓸 평문으로 변환."""
    if not ai_summary:
        return ""
    if isinstance(ai_summary, str):
        try:
            ai_summary = json.loads(ai_summary)
        except (TypeError, ValueError):
            return ai_summary
    if not isinstance(ai_summary, dict):
        return str(ai_summary)
    parts = []
    title = ai_summary.get("title")
    if title:
        parts.append(f"[{title}]")
    summary = ai_summary.get("summary")
    if isinstance(summary, dict):
        for k, v in summary.items():
            if v and str(v).strip():
                parts.append(f"- {k}: {v}")
    elif summary:
        parts.append(str(summary))
    return "\n".join(parts)


async def _collect_items_from_memos(
    db: AsyncSession, memo_ids: list[int], user_id: int
) -> list[dict]:
    if not memo_ids:
        return []
    query = (
        select(VisitMemo)
        .options(selectinload(VisitMemo.doctor).selectinload(Doctor.hospital))
        .where(
            VisitMemo.id.in_(memo_ids),
            VisitMemo.user_id == user_id,
        )
    )
    rows = (await db.execute(query)).scalars().all()
    rows.sort(key=lambda m: m.visit_date or datetime.min)
    items = []
    for m in rows:
        doc = m.doctor
        hosp = doc.hospital if doc else None
        items.append({
            "visit_date": m.visit_date.strftime("%Y-%m-%d %H:%M") if m.visit_date else None,
            "doctor_name": doc.name if doc else (m.doctor_name_snapshot or ""),
            "hospital_name": hosp.name if hosp else (m.hospital_name_snapshot or ""),
            "department": doc.department if doc else (m.doctor_dept_snapshot or ""),
            "title": m.title,
            "raw_memo": m.raw_memo,
            "ai_summary_text": _ai_summary_to_text(_parse_ai(m.ai_summary)),
        })
    return items


async def _collect_items_from_reports(
    db: AsyncSession, report_ids: list[int], user_id: int
) -> list[dict]:
    if not report_ids:
        return []
    query = select(Report).where(
        Report.id.in_(report_ids),
        Report.user_id == user_id,
    )
    rows = (await db.execute(query)).scalars().all()
    rows.sort(key=lambda r: r.period_start)
    items = []
    for r in rows:
        items.append({
            "visit_date": r.period_start,
            "doctor_name": "",
            "hospital_name": "",
            "department": "",
            "title": r.title or f"{r.report_type} 보고서",
            "raw_memo": r.raw_combined or "",
            "ai_summary_text": _ai_summary_to_text(_parse_ai(r.ai_summary)),
        })
    return items


def _items_to_raw_combined(items: list[dict]) -> str:
    blocks = []
    for i, it in enumerate(items, 1):
        meta = " · ".join(
            v for v in [it.get("visit_date"), it.get("doctor_name"), it.get("hospital_name")]
            if v
        )
        body = it.get("ai_summary_text") or it.get("raw_memo") or ""
        blocks.append(f"[{i}] {meta}\n{body}")
    return "\n\n".join(blocks)


def _period_label(period_start: str, period_end: str, report_type: str) -> str:
    if report_type == "daily" or period_start == period_end:
        return period_start
    return f"{period_start} ~ {period_end}"


async def _load_template_settings(
    db: AsyncSession, template_id: Optional[int]
) -> tuple[Optional[str], Optional[list[str]]]:
    """템플릿 id 로 prompt_addon 과 custom_fields 추출. 템플릿 없으면 (None, None)."""
    if not template_id:
        return None, None
    t = (await db.execute(
        select(MemoTemplate).where(MemoTemplate.id == template_id)
    )).scalar_one_or_none()
    if not t:
        return None, None
    custom_fields: Optional[list[str]] = None
    try:
        parsed = json.loads(t.fields) if t.fields else None
        if isinstance(parsed, list) and parsed:
            custom_fields = [str(f) for f in parsed if str(f).strip()]
    except (TypeError, ValueError):
        pass
    return t.prompt_addon, custom_fields


@router.post("", summary="보고서 생성 (AI 종합)")
async def create_report(
    payload: ReportCreate,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    if payload.report_type not in ("daily", "weekly"):
        raise HTTPException(status_code=400, detail="report_type 은 'daily' 또는 'weekly' 여야 합니다.")
    if not payload.memo_ids and not payload.report_ids:
        raise HTTPException(status_code=400, detail="memo_ids 또는 report_ids 중 하나는 필요합니다.")
    if payload.memo_ids and payload.report_ids:
        raise HTTPException(status_code=400, detail="memo_ids 와 report_ids 를 동시에 지정할 수 없습니다.")

    if payload.memo_ids:
        items = await _collect_items_from_memos(db, payload.memo_ids, user.id)
    else:
        items = await _collect_items_from_reports(db, payload.report_ids, user.id)

    if not items:
        raise HTTPException(status_code=404, detail="해당 ID로 종합할 메모/보고서를 찾지 못했습니다.")

    # 템플릿 fields/prompt_addon 활용 (옵션)
    prompt_addon, custom_fields = await _load_template_settings(db, payload.template_id)

    period_label = _period_label(payload.period_start, payload.period_end, payload.report_type)
    ai_result = await summarize_report(
        items=items,
        report_type=payload.report_type,
        period_label=period_label,
        prompt_addon=prompt_addon,
        custom_fields=custom_fields,
    )

    report = Report(
        user_id=user.id,
        report_type=payload.report_type,
        period_start=payload.period_start,
        period_end=payload.period_end,
        title=payload.title or ai_result.get("title") or f"{period_label} {payload.report_type} 보고서",
        source_memo_ids=json.dumps(payload.memo_ids) if payload.memo_ids else None,
        source_report_ids=json.dumps(payload.report_ids) if payload.report_ids else None,
        raw_combined=_items_to_raw_combined(items),
        ai_summary=json.dumps(ai_result, ensure_ascii=False),
        template_id=payload.template_id,
    )
    db.add(report)
    await db.commit()
    await db.refresh(report)
    return _serialize(report)


@router.get("", summary="보고서 목록")
async def list_reports(
    from_date: Optional[str] = Query(None, alias="from"),
    to_date: Optional[str] = Query(None, alias="to"),
    report_type: Optional[str] = None,
    limit: int = Query(50, ge=1, le=200),
    offset: int = Query(0, ge=0),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    query = select(Report).where(Report.user_id == user.id)
    if report_type in ("daily", "weekly"):
        query = query.where(Report.report_type == report_type)
    if from_date:
        query = query.where(Report.period_end >= from_date)
    if to_date:
        query = query.where(Report.period_start <= to_date)
    query = query.order_by(Report.period_start.desc(), Report.id.desc()).offset(offset).limit(limit)
    rows = (await db.execute(query)).scalars().all()
    return [_serialize(r) for r in rows]


@router.get("/{report_id}", summary="보고서 상세")
async def get_report(
    report_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    r = (await db.execute(
        select(Report).where(Report.id == report_id, Report.user_id == user.id)
    )).scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    return _serialize(r)


@router.post("/{report_id}/regenerate", summary="AI 재정리")
async def regenerate_report(
    report_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    r = (await db.execute(
        select(Report).where(Report.id == report_id, Report.user_id == user.id)
    )).scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")

    memo_ids = _parse_json_array(r.source_memo_ids)
    report_ids = _parse_json_array(r.source_report_ids)
    if memo_ids:
        items = await _collect_items_from_memos(db, memo_ids, user.id)
    elif report_ids:
        items = await _collect_items_from_reports(db, report_ids, user.id)
    else:
        raise HTTPException(status_code=400, detail="원본 메모/보고서 정보가 없습니다.")

    prompt_addon, custom_fields = await _load_template_settings(db, r.template_id)

    period_label = _period_label(r.period_start, r.period_end, r.report_type)
    ai_result = await summarize_report(
        items=items,
        report_type=r.report_type,
        period_label=period_label,
        prompt_addon=prompt_addon,
        custom_fields=custom_fields,
    )

    r.ai_summary = json.dumps(ai_result, ensure_ascii=False)
    r.raw_combined = _items_to_raw_combined(items)
    if not r.title:
        r.title = ai_result.get("title")
    await db.commit()
    await db.refresh(r)
    return _serialize(r)


@router.delete("/{report_id}", summary="보고서 삭제")
async def delete_report(
    report_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    r = (await db.execute(
        select(Report).where(Report.id == report_id, Report.user_id == user.id)
    )).scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")
    await db.delete(r)
    await db.commit()
    return {"status": "deleted", "id": report_id}


# ─────────── DOCX 다운로드 ───────────

def _build_docx(report: Report) -> bytes:
    """보고서를 docx 바이트로 변환.

    한글 깨짐 방지: python-docx 디폴트 폰트는 Calibri (서양 전용) 라 일부 Word
    환경에서 한글이 eastAsia 슬롯 폰트 매핑 실패로 □□□ 로 표시된다. 모든 run
    에 '맑은 고딕' + eastAsia 슬롯을 명시.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail="python-docx 미설치. backend 에서 `pip install python-docx` 후 재시작하세요.",
        ) from e

    KO_FONT = "맑은 고딕"

    def _apply_font(run):
        run.font.name = KO_FONT
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), KO_FONT)
        rFonts.set(qn("w:hAnsi"), KO_FONT)
        rFonts.set(qn("w:eastAsia"), KO_FONT)
        rFonts.set(qn("w:cs"), KO_FONT)

    def _add_para(text: str = ""):
        p = doc.add_paragraph()
        if text:
            r = p.add_run(text)
            _apply_font(r)
        return p

    def _add_heading(text: str, level: int):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            _apply_font(r)
        return h

    doc = Document()

    # Normal 스타일 자체도 한글 폰트로 (테마 fallback 차단)
    normal = doc.styles["Normal"]
    normal.font.name = KO_FONT
    n_rPr = normal.element.get_or_add_rPr()
    n_rFonts = n_rPr.find(qn("w:rFonts"))
    if n_rFonts is None:
        n_rFonts = OxmlElement("w:rFonts")
        n_rPr.append(n_rFonts)
    n_rFonts.set(qn("w:ascii"), KO_FONT)
    n_rFonts.set(qn("w:hAnsi"), KO_FONT)
    n_rFonts.set(qn("w:eastAsia"), KO_FONT)

    type_label = "주간" if report.report_type == "weekly" else "일일"
    period_label = (
        report.period_start
        if report.period_start == report.period_end
        else f"{report.period_start} ~ {report.period_end}"
    )

    # 제목
    h = _add_heading(report.title or f"{type_label} 보고서", level=1)
    if h.runs:
        h.runs[0].font.size = Pt(18)

    # 메타 정보 줄
    meta = doc.add_paragraph()
    r1 = meta.add_run(f"{type_label} 보고서  |  기간: {period_label}")
    r1.italic = True
    _apply_font(r1)
    if report.created_at:
        r2 = meta.add_run(f"  |  작성일: {report.created_at.strftime('%Y-%m-%d %H:%M')}")
        r2.italic = True
        _apply_font(r2)

    # AI 종합 섹션
    summary = None
    if report.ai_summary:
        try:
            summary = json.loads(report.ai_summary)
        except (TypeError, ValueError):
            summary = None
    sections = (summary or {}).get("summary") if isinstance(summary, dict) else None

    if isinstance(sections, dict) and sections:
        _add_para()  # 빈 줄
        for key, value in sections.items():
            if not value or not str(value).strip():
                continue
            _add_heading(str(key), level=2)
            for line in str(value).split("\n"):
                _add_para(line)
    else:
        _add_para()
        _add_para("(AI 정리 결과 없음)")

    # 원본(감사용) — 페이지 나눔 후 부록으로
    if report.raw_combined:
        doc.add_page_break()
        _add_heading("부록 — 원본 메모/기록", level=2)
        for line in report.raw_combined.split("\n"):
            _add_para(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/{report_id}/docx", summary="보고서 docx 다운로드")
async def download_report_docx(
    report_id: int,
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    r = (await db.execute(
        select(Report).where(Report.id == report_id, Report.user_id == user.id)
    )).scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="보고서를 찾을 수 없습니다.")

    data = _build_docx(r)
    type_label = "주간" if r.report_type == "weekly" else "일일"
    suffix = r.period_start if r.period_start == r.period_end else f"{r.period_start}_{r.period_end}"
    filename = f"{type_label}보고서_{suffix}.docx"
    # RFC 5987: 한글 파일명을 위해 filename* 사용
    from urllib.parse import quote
    encoded_name = quote(filename)
    headers = {
        "Content-Disposition": f"attachment; filename=\"report.docx\"; filename*=UTF-8''{encoded_name}",
    }
    return StreamingResponse(
        io.BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
